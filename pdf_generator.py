import os
import glob
import img2pdf
import re
from utils import sanitize_filename, parse_image_timestamp
from transcript import clean_transcript_text

def convert_screenshots_to_pdf(images_folder, output_folder, video_name):
    """Convert screenshots to PDF"""
    safe_name = sanitize_filename(video_name)
    output_pdf_path = os.path.join(output_folder, f"{safe_name}.pdf")
    
    # Get all PNG files sorted
    image_files = sorted(glob.glob(os.path.join(images_folder, "*.png")))
    
    if not image_files:
        print("No images found to convert to PDF")
        return None
    
    print(f'Converting {len(image_files)} images to PDF...')
    
    try:
        with open(output_pdf_path, "wb") as f:
            f.write(img2pdf.convert(image_files))
        print(f'PDF created successfully: {output_pdf_path}')
        return output_pdf_path
    except Exception as e:
        print(f'Error creating PDF: {e}')
        raise


def sync_images_with_transcript(images_folder, transcript_file, output_folder, video_name="combined"):
    """Sync images with transcript and create a combined document"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        
        print("\nSyncing images with transcript and creating combined document...")
        
        # Parse transcript
        transcript_entries = []
        with open(transcript_file, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                # Extract timestamp and text: [00:00:01] text
                match = re.match(r'\[(\d{2}:\d{2}:\d{2})\]\s*(.+)', line)
                if match:
                    timestamp = match.group(1)
                    text = match.group(2)
                    transcript_entries.append((timestamp, text))
        
        # Get all images with their timestamps
        image_files = sorted(glob.glob(os.path.join(images_folder, "*.png")))
        image_data = []
        for img_file in image_files:
            timestamp = parse_image_timestamp(os.path.basename(img_file))
            if timestamp:
                image_data.append((timestamp, img_file))
        
        if not image_data:
            print("No images found to sync")
            return None
        
        # Create PDF
        safe_name = sanitize_filename(video_name)
        output_pdf = os.path.join(output_folder, f"{safe_name}.pdf")
        doc = SimpleDocTemplate(output_pdf, pagesize=letter)
        story = []
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor='#333333',
            spaceAfter=12,
            alignment=TA_CENTER
        )
        timestamp_style = ParagraphStyle(
            'Timestamp',
            parent=styles['Normal'],
            fontSize=10,
            textColor='#666666',
            spaceAfter=6,
            alignment=TA_LEFT
        )
        text_style = ParagraphStyle(
            'TranscriptText',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            spaceAfter=12,
            alignment=TA_LEFT
        )
        
        # Match images with transcript
        image_idx = 0
        transcript_idx = 0
        
        while image_idx < len(image_data):
            img_timestamp, img_path = image_data[image_idx]
            
            # Find transcript entries that match this image's timestamp
            # Look for transcript entries within 30 seconds of the image
            matching_transcript = []
            
            # Convert timestamps to seconds for comparison
            def timestamp_to_seconds(ts):
                parts = ts.split(':')
                return int(parts[0]) * 3600 + int(parts[1]) * 60 + int(parts[2])
            
            img_seconds = timestamp_to_seconds(img_timestamp)
            
            # Find transcript entries near this timestamp
            for ts, text in transcript_entries:
                ts_seconds = timestamp_to_seconds(ts)
                if abs(ts_seconds - img_seconds) <= 30:  # Within 30 seconds
                    matching_transcript.append((ts, text))
            
            # Add image (Optimized)
            try:
                # Import PIL if not already imported (it is imported as Image from reportlab, so use PILImage)
                from PIL import Image as PILImage
                import io
                
                # Open image and convert to JPEG in memory
                with PILImage.open(img_path) as pil_img:
                    # Convert to RGB if needed
                    if pil_img.mode in ('RGBA', 'LA', 'P'):
                        rgb_img = PILImage.new('RGB', pil_img.size, (255, 255, 255))
                        if pil_img.mode == 'P':
                            pil_img = pil_img.convert('RGBA')
                        rgb_img.paste(pil_img, mask=pil_img.split()[-1] if pil_img.mode == 'RGBA' else None)
                        pil_img = rgb_img
                    
                    # Resize if too large (e.g. > 1600px width) to reduce size
                    max_width = 1600
                    if pil_img.width > max_width:
                        ratio = max_width / pil_img.width
                        new_height = int(pil_img.height * ratio)
                        pil_img = pil_img.resize((max_width, new_height), PILImage.LANCZOS)

                    # Save to BytesIO as JPEG
                    img_byte_arr = io.BytesIO()
                    pil_img.save(img_byte_arr, format='JPEG', quality=60, optimize=True)
                    img_byte_arr.seek(0)
                    
                    # Create ReportLab Image from BytesIO
                    # We need to use a temporary file because ReportLab Image prefers paths, 
                    # but it can accept a file-like object if we are careful.
                    # Actually, ReportLab's Image class accepts a filename or a file-like object.
                    
                    img = Image(img_byte_arr, width=6*inch, height=4.5*inch)
                    img.drawHeight = 4.5*inch
                    img.drawWidth = 6*inch
                    story.append(img)
                    story.append(Spacer(1, 0.2*inch))
            except Exception as e:
                print(f"Warning: Could not add image {img_path}: {e}")
                # Fallback to original
                try:
                    img = Image(img_path, width=6*inch, height=4.5*inch)
                    story.append(img)
                    story.append(Spacer(1, 0.2*inch))
                except:
                    pass
            
            # Add matching transcript text
            if matching_transcript:
                # Clean and format transcript
                transcript_texts = [text for _, text in matching_transcript]
                combined_text = ' '.join(transcript_texts)
                cleaned_text = clean_transcript_text(combined_text)
                
                if cleaned_text:
                    story.append(Paragraph(f"<b>Transcript ({img_timestamp}):</b>", timestamp_style))
                    story.append(Spacer(1, 0.1*inch))
                    story.append(Paragraph(cleaned_text.replace('\n', '<br/>'), text_style))
                    story.append(Spacer(1, 0.3*inch))
            
            story.append(PageBreak())
            image_idx += 1
        
        # Build PDF
        doc.build(story)
        print(f"Combined PDF created: {output_pdf}")
        return output_pdf
        
    except ImportError as e:
        print(f"Error: Missing required library. Please install: pip install reportlab pyspellchecker")
        print(f"Error details: {e}")
        return None
    except Exception as e:
        print(f"Error creating combined document: {e}")
        import traceback
        traceback.print_exc()
        return None

def sync_images_with_transcript_docx(images_folder, transcript_file, output_folder, video_name="combined"):
    """Sync images with transcript and create a combined DOCX document"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        print("\nSyncing images with transcript and creating combined DOCX document...")
        
        # Parse transcript
        transcript_entries = []
        
        # Try different encodings
        content = None
        encodings = ['utf-8', 'utf-16', 'utf-16-le', 'utf-16-be', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(transcript_file, 'r', encoding=encoding) as f:
                    content = f.read()
                print(f"Successfully read transcript with encoding: {encoding}")
                break
            except UnicodeDecodeError:
                continue
            except Exception as e:
                print(f"Error reading transcript with {encoding}: {e}")
                continue
                
        if content is None:
            print(f"Failed to read transcript file {transcript_file} with any of the attempted encodings.")
            return None
            
        for line in content.splitlines():
            line = line.strip()
            if not line:
                continue
            # Extract timestamp and text: [00:00:01] text
            match = re.match(r'\[(\d{2}:\d{2}:\d{2})\]\s*(.+)', line)
            if match:
                timestamp = match.group(1)
                text = match.group(2)
                transcript_entries.append((timestamp, text))
        
        # Get all images with their timestamps
        image_files = sorted(glob.glob(os.path.join(images_folder, "*.png")))
        image_data = []
        for img_file in image_files:
            timestamp = parse_image_timestamp(os.path.basename(img_file))
            if timestamp:
                image_data.append((timestamp, img_file))
        
        if not image_data:
            print("No images found to sync")
            return None
        
        # Create DOCX
        safe_name = sanitize_filename(video_name)
        output_docx = os.path.join(output_folder, f"{safe_name}.docx")
        doc = Document()
        
        # Match images with transcript
        image_idx = 0
        
        # Import PIL for image optimization
        from PIL import Image as PILImage
        import io
        
        while image_idx < len(image_data):
            img_timestamp, img_path = image_data[image_idx]
            
            # Find transcript entries that match this image's timestamp
            # Look for transcript entries within 30 seconds of the image
            matching_transcript = []
            
            # Convert timestamps to seconds for comparison
            def timestamp_to_seconds(ts):
                parts = ts.split(':')
                return int(parts[0]) * 3600 + int(parts[1]) * 60 + int(parts[2])
            
            img_seconds = timestamp_to_seconds(img_timestamp)
            
            # Find transcript entries near this timestamp
            for ts, text in transcript_entries:
                ts_seconds = timestamp_to_seconds(ts)
                if abs(ts_seconds - img_seconds) <= 30:  # Within 30 seconds
                    matching_transcript.append((ts, text))
            
            # Add image (Optimized)
            try:
                # Open image and convert to JPEG in memory
                with PILImage.open(img_path) as img:
                    # Convert to RGB if needed
                    if img.mode in ('RGBA', 'LA', 'P'):
                        rgb_img = PILImage.new('RGB', img.size, (255, 255, 255))
                        if img.mode == 'P':
                            img = img.convert('RGBA')
                        rgb_img.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                        img = rgb_img
                    
                    # Resize if too large (e.g. > 1600px width) to reduce size
                    max_width = 1600
                    if img.width > max_width:
                        ratio = max_width / img.width
                        new_height = int(img.height * ratio)
                        img = img.resize((max_width, new_height), PILImage.LANCZOS)

                    # Save to BytesIO as JPEG
                    img_byte_arr = io.BytesIO()
                    img.save(img_byte_arr, format='JPEG', quality=60, optimize=True)
                    img_byte_arr.seek(0)
                    
                    doc.add_picture(img_byte_arr, width=Inches(6))
            except Exception as e:
                print(f"Warning: Could not add image {img_path}: {e}")
            
            # Add matching transcript text
            if matching_transcript:
                # Clean and format transcript
                transcript_texts = [text for _, text in matching_transcript]
                combined_text = ' '.join(transcript_texts)
                cleaned_text = clean_transcript_text(combined_text)
                
                if cleaned_text:
                    p = doc.add_paragraph()
                    run = p.add_run(f"Transcript ({img_timestamp}):")
                    run.bold = True
                    doc.add_paragraph(cleaned_text)
            
            # Add page break if not the last image
            if image_idx < len(image_data) - 1:
                doc.add_page_break()
                
            image_idx += 1
        
        # Save DOCX
        doc.save(output_docx)
        print(f"Combined DOCX created: {output_docx}")
        return output_docx
        
    except ImportError as e:
        print(f"Error: Missing required library. Please install: pip install python-docx")
        print(f"Error details: {e}")
        return None
    except Exception as e:
        print(f"Error creating combined DOCX document: {e}")
        import traceback
        traceback.print_exc()
        return None

def create_pdf_from_data(slides_data, output_pdf_path):
    """
    Create a PDF from a list of slide data objects.
    slides_data: list of dicts { 'image_path': str, 'text': str, 'timestamp': str }
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT
        from PIL import Image as PILImage
        import io

        doc = SimpleDocTemplate(output_pdf_path, pagesize=letter)
        story = []
        
        styles = getSampleStyleSheet()
        timestamp_style = ParagraphStyle(
            'Timestamp',
            parent=styles['Normal'],
            fontSize=10,
            textColor='#666666',
            spaceAfter=6,
            alignment=TA_LEFT
        )
        text_style = ParagraphStyle(
            'TranscriptText',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            spaceAfter=12,
            alignment=TA_LEFT
        )

        for slide in slides_data:
            img_path = slide['image_path']
            text = slide['text']
            timestamp = slide.get('timestamp', '')

            # Add Image
            try:
                with PILImage.open(img_path) as pil_img:
                    # Convert to RGB if needed
                    if pil_img.mode in ('RGBA', 'LA', 'P'):
                        rgb_img = PILImage.new('RGB', pil_img.size, (255, 255, 255))
                        if pil_img.mode == 'P':
                            pil_img = pil_img.convert('RGBA')
                        rgb_img.paste(pil_img, mask=pil_img.split()[-1] if pil_img.mode == 'RGBA' else None)
                        pil_img = rgb_img
                    
                    # Resize/Compress
                    max_width = 1600
                    if pil_img.width > max_width:
                        ratio = max_width / pil_img.width
                        new_height = int(pil_img.height * ratio)
                        pil_img = pil_img.resize((max_width, new_height), PILImage.LANCZOS)

                    img_byte_arr = io.BytesIO()
                    pil_img.save(img_byte_arr, format='JPEG', quality=60, optimize=True)
                    img_byte_arr.seek(0)
                    
                    img = Image(img_byte_arr, width=6*inch, height=4.5*inch)
                    story.append(img)
                    story.append(Spacer(1, 0.2*inch))
            except Exception as e:
                print(f"Warning: Could not add image {img_path}: {e}")

            # Add Text
            if text:
                if timestamp:
                    story.append(Paragraph(f"<b>Transcript ({timestamp}):</b>", timestamp_style))
                story.append(Spacer(1, 0.1*inch))
                # Handle newlines
                formatted_text = text.replace('\n', '<br/>')
                story.append(Paragraph(formatted_text, text_style))
                story.append(Spacer(1, 0.3*inch))
            
            story.append(PageBreak())

        doc.build(story)
        return output_pdf_path

    except Exception as e:
        print(f"Error creating PDF from data: {e}")
        import traceback
        traceback.print_exc()
        return None

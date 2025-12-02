
from docx import Document
from docx.shared import Pt
import re
import os

INPUT_FILE = r"C:\Users\vinay\video2pdf\output\Bootcamp Classroom - Week 7 Day 1 - 9_30 trade\generated_notes.md"
OUTPUT_FILE = r"C:\Users\vinay\video2pdf\output\Bootcamp Classroom - Week 7 Day 1 - 9_30 trade\generated_notes_full.docx"

def md_to_docx(md_file, docx_file):
    if not os.path.exists(md_file):
        print(f"File not found: {md_file}")
        return

    with open(md_file, 'r', encoding='utf-8') as f:
        text = f.read()

    doc = Document()
    
    # Simple parser for the specific format we generated
    # ## [Slide Title] -> Heading 2
    # [[IMAGE_PATH: ...]] -> Image
    # [[IMAGE_RELEVANCE: ...]] -> Caption
    # * **Key Point**: Desc -> Bullet point with bold start
    # --- -> Page Break
    
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('---'):
            doc.add_page_break()
        elif line.startswith('[[IMAGE_PATH: '):
            image_path = line[14:-2]
            if os.path.exists(image_path):
                try:
                    # Add picture, width 6 inches
                    doc.add_picture(image_path, width=Pt(432)) # 6 inches * 72 points
                except Exception as e:
                    print(f"Error adding image {image_path}: {e}")
                    doc.add_paragraph(f"[Image: {os.path.basename(image_path)}]")
        elif line.startswith('[[IMAGE_RELEVANCE: '):
            relevance = line[19:-2]
            p = doc.add_paragraph()
            run = p.add_run(f"Relevance: {relevance}")
            run.italic = True
            run.font.size = Pt(9)
        elif line.startswith('* '):
            # Bullet point
            p = doc.add_paragraph(style='List Bullet')
            content = line[2:]
            
            # Handle bolding **text**
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        else:
            doc.add_paragraph(line)

    doc.save(docx_file)
    print(f"DOCX generated: {docx_file}")

if __name__ == "__main__":
    md_to_docx(INPUT_FILE, OUTPUT_FILE)

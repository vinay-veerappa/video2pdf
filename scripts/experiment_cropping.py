import os
import glob
import cv2
import numpy as np
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Inches

def get_content_bbox(img, threshold_ratio=0.1):
    """
    Find crop coordinates using projection profiles (row/col variance).
    Returns (x, y, w, h)
    """
    if len(img.shape) == 3:
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    else:
        gray = img

    # Calculate variance along rows and columns using Sobel
    sobelx = cv2.Sobel(gray, cv2.CV_64F, 1, 0, ksize=3)
    sobely = cv2.Sobel(gray, cv2.CV_64F, 0, 1, ksize=3)
    magnitude = np.sqrt(sobelx**2 + sobely**2)
    
    # Project onto axes
    row_projection = np.sum(magnitude, axis=1)
    col_projection = np.sum(magnitude, axis=0)
    
    # Normalize
    row_max = np.max(row_projection)
    col_max = np.max(col_projection)
    
    if row_max == 0 or col_max == 0:
        return 0, 0, img.shape[1], img.shape[0]
        
    row_projection = row_projection / row_max
    col_projection = col_projection / col_max
    
    # Threshold
    active_rows = np.where(row_projection > threshold_ratio)[0]
    active_cols = np.where(col_projection > threshold_ratio)[0]
    
    if len(active_rows) == 0 or len(active_cols) == 0:
        return 0, 0, img.shape[1], img.shape[0]
        
    y_min, y_max = active_rows[0], active_rows[-1]
    x_min, x_max = active_cols[0], active_cols[-1]
    
    # Refinement: Move inwards to avoid border edges
    h, w = img.shape[:2]
    margin_x = int(w * 0.01)
    margin_y = int(h * 0.01)
    
    x_min = min(x_min + margin_x, w)
    y_min = min(y_min + margin_y, h)
    x_max = max(x_max - margin_x, 0)
    y_max = max(y_max - margin_y, 0)
    
    # Ensure valid coordinates
    if x_max <= x_min or y_max <= y_min:
        return 0, 0, w, h
        
    return x_min, y_min, x_max - x_min, y_max - y_min

def run_experiment(input_dir, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Get all png images
    images = glob.glob(os.path.join(input_dir, "*.png"))
    
    # Select 3 diverse images
    if len(images) < 3:
        selected_images = images
    else:
        images.sort()
        selected_images = [images[0], images[len(images)//2], images[-1]]

    print(f"Selected images: {[os.path.basename(i) for i in selected_images]}")

    # Initialize DOCX
    doc = Document()
    doc.add_heading('Image Cropping Experiment', 0)

    html_content = """
    <html>
    <head>
        <style>
            body { font-family: sans-serif; padding: 20px; }
            table { border-collapse: collapse; width: 100%; }
            th, td { border: 1px solid #ddd; padding: 8px; text-align: center; vertical-align: top; }
            th { background-color: #f2f2f2; }
            img { max-width: 400px; height: auto; display: block; margin: 0 auto; border: 1px solid #ccc; }
            .meta { font-size: 12px; color: #555; margin-top: 5px; }
        </style>
    </head>
    <body>
        <h1>Image Cropping Experiment</h1>
        <table>
            <tr>
                <th>Original</th>
                <th>Detected Box</th>
                <th>Auto-Cropped</th>
            </tr>
    """

    for img_path in selected_images:
        basename = os.path.basename(img_path)
        
        # Read with OpenCV
        img = cv2.imread(img_path)
        if img is None: continue
        
        h, w = img.shape[:2]
        
        # Calculate crop
        x, y, cw, ch = get_content_bbox(img)
        
        # Create visualization (Draw rectangle on copy)
        vis_img = img.copy()
        cv2.rectangle(vis_img, (x, y), (x+cw, y+ch), (0, 0, 255), 2)
        
        # Perform crop
        cropped_img = img[y:y+ch, x:x+cw]
        
        # Save files
        vis_path = os.path.join(output_dir, f"vis_{basename}")
        crop_path = os.path.join(output_dir, f"crop_{basename}")
        
        cv2.imwrite(vis_path, vis_img)
        cv2.imwrite(crop_path, cropped_img)
        
        # Add to HTML
        html_content += f"""
        <tr>
            <td>
                <img src='file:///{img_path.replace(os.sep, '/')}' onclick="window.open(this.src, '_blank');">
                <div class='meta'>{basename}<br>{w}x{h}</div>
            </td>
            <td>
                <img src='file:///{vis_path.replace(os.sep, '/')}' onclick="window.open(this.src, '_blank');">
                <div class='meta'>Box: x={x}, y={y}, w={cw}, h={ch}</div>
            </td>
            <td>
                <img src='file:///{crop_path.replace(os.sep, '/')}' onclick="window.open(this.src, '_blank');">
                <div class='meta'>Result<br>{cw}x{ch}</div>
            </td>
        </tr>
        """
        
        # Add to DOCX
        doc.add_heading(f'Image: {basename}', level=1)
        
        doc.add_heading('Original', level=2)
        doc.add_picture(img_path, width=Inches(5))
        
        doc.add_heading('Auto-Cropped', level=2)
        doc.add_picture(crop_path, width=Inches(5))
        
        doc.add_paragraph(f"Original Size: {w}x{h}")
        doc.add_paragraph(f"Cropped Size: {cw}x{ch}")
        doc.add_paragraph(f"Crop Box: x={x}, y={y}, w={cw}, h={ch}")
        doc.add_paragraph("-" * 50)

    html_content += """
        </table>
    </body>
    </html>
    """

    # Save HTML
    with open(os.path.join(output_dir, "cropping_experiment.html"), "w") as f:
        f.write(html_content)

    # Save DOCX
    doc.save(os.path.join(output_dir, "cropping_experiment.docx"))
    
    print(f"Experiment complete. Results saved to {output_dir}")

if __name__ == "__main__":
    input_dir = r"C:\Users\vinay\video2pdf\output\Bootcamp Classroom - Week 4 Day 1 - Candle Stick Science intro\images"
    output_dir = r"C:\Users\vinay\video2pdf\cropping_experiment"
    run_experiment(input_dir, output_dir)

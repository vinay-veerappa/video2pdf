import os
import glob
from PIL import Image
from docx import Document
from docx.shared import Inches
import shutil

def get_size_format(b, factor=1024, suffix="B"):
    """
    Scale bytes to its proper byte format
    e.g:
        1253656 => '1.20MB'
        1253656678 => '1.17GB'
    """
    for unit in ["", "K", "M", "G", "T", "P", "E", "Z"]:
        if b < factor:
            return f"{b:.2f}{unit}{suffix}"
        b /= factor
    return f"{b:.2f}Y{suffix}"

def run_experiment(input_dir, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Get all png images
    images = glob.glob(os.path.join(input_dir, "*.png"))
    
    # Select 3 diverse images (small, medium, large if possible, or just first, middle, last)
    if len(images) < 3:
        selected_images = images
    else:
        images.sort()
        selected_images = [images[0], images[len(images)//2], images[-1]]

    print(f"Selected images: {[os.path.basename(i) for i in selected_images]}")

    results = []
    
    # Qualities to test
    qualities = [70, 60, 50, 40, 30]

    # Initialize DOCX
    doc = Document()
    doc.add_heading('Image Compression Experiment', 0)

    html_content = """
    <html>
    <head>
        <style>
            body { font-family: sans-serif; padding: 20px; }
            table { border-collapse: collapse; width: 100%; }
            th, td { border: 1px solid #ddd; padding: 8px; text-align: center; vertical-align: top; }
            th { background-color: #f2f2f2; }
            img { max-width: 300px; height: auto; display: block; margin: 0 auto; }
            .meta { font-size: 12px; color: #555; margin-top: 5px; }
            .reduction { color: green; font-weight: bold; }
        </style>
    </head>
    <body>
        <h1>Image Compression Experiment</h1>
        <table>
            <tr>
                <th>Original</th>
    """
    
    for q in qualities:
        html_content += f"<th>Quality {q}</th>"
    html_content += "</tr>"

    for img_path in selected_images:
        basename = os.path.basename(img_path)
        original_size = os.path.getsize(img_path)
        
        # Load image
        with Image.open(img_path) as img:
            # Convert to RGB for JPEG
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
            
            row_html = f"<tr><td><img src='file:///{img_path.replace(os.sep, '/')}'><div class='meta'>{basename}<br>{get_size_format(original_size)}</div></td>"
            
            # Add to DOCX
            doc.add_heading(f'Image: {basename}', level=1)
            doc.add_paragraph(f'Original Size: {get_size_format(original_size)}')
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph("This is the original PNG image for reference.")
            doc.add_paragraph("-" * 50)

            for q in qualities:
                # Save compressed
                out_name = f"{os.path.splitext(basename)[0]}_q{q}.jpg"
                out_path = os.path.join(output_dir, out_name)
                img.save(out_path, "JPEG", quality=q, optimize=True)
                
                compressed_size = os.path.getsize(out_path)
                reduction = (original_size - compressed_size) / original_size * 100
                
                row_html += f"""
                <td>
                    <img src='file:///{out_path.replace(os.sep, '/')}' onclick="window.open(this.src, '_blank');">
                    <div class='meta'>
                        Size: {get_size_format(compressed_size)}<br>
                        <span class='reduction'>-{reduction:.1f}%</span>
                    </div>
                </td>
                """
                
                # Add to DOCX
                doc.add_heading(f'Quality: {q}', level=2)
                doc.add_paragraph(f'Compressed Size: {get_size_format(compressed_size)} (Reduction: {reduction:.1f}%)')
                doc.add_picture(out_path, width=Inches(6))
                doc.add_paragraph(f"This image was compressed with JPEG quality {q}. Notice any artifacts around text or sharp edges.")
                doc.add_paragraph("Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.")
                doc.add_paragraph("-" * 50)

            row_html += "</tr>"
            html_content += row_html

    html_content += """
        </table>
    </body>
    </html>
    """

    # Save HTML
    with open(os.path.join(output_dir, "compression_experiment.html"), "w") as f:
        f.write(html_content)

    # Save DOCX
    doc.save(os.path.join(output_dir, "compression_experiment.docx"))
    
    print(f"Experiment complete. Results saved to {output_dir}")

if __name__ == "__main__":
    input_dir = r"C:\Users\vinay\video2pdf\output\Bootcamp Classroom - Week 4 Day 1 - Candle Stick Science intro\images"
    output_dir = r"C:\Users\vinay\video2pdf\compression_experiment"
    run_experiment(input_dir, output_dir)
